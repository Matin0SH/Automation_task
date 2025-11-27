"""
File extractors for different document formats (PDF, DOCX)

Upgraded to preserve lists and tables:
- DOCX: prefixes list-style paragraphs, serializes tables row-by-row
- PDF: uses pdfplumber for better text flow + table extraction
"""

import os
from pathlib import Path
from docx import Document

try:
    import pdfplumber
except ImportError as e:
    raise ImportError("pdfplumber is required for PDF extraction. Install with: pip install pdfplumber") from e


class DocumentExtractor:
    """Base class for document extraction"""

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from a document based on its file extension

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text as string

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.docx':
            return DocumentExtractor._extract_docx(file_path)
        elif file_ext == '.pdf':
            return DocumentExtractor._extract_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX, keeping lists and tables."""
        try:
            doc = Document(file_path)
            chunks = []

            # Paragraphs (basic list detection via style name)
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
                is_list = any(key in style_name for key in ["list", "bullet", "number"])
                prefix = "- " if is_list else ""
                chunks.append(f"{prefix}{text}")

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    chunks.append(" | ".join(cells))
                chunks.append("")  # blank line after each table

            return "\n".join([c for c in chunks if c is not None]).strip()
        except Exception as e:
            raise Exception(f"Error extracting DOCX file {file_path}: {str(e)}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text and tables from PDF using pdfplumber."""
        try:
            chunks = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Page text
                    text = (page.extract_text() or "").strip()
                    if text:
                        chunks.append(text)

                    # Tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            cells = [(cell or "").strip() for cell in row]
                            chunks.append(" | ".join(cells))
                        chunks.append("")  # blank line after each table

            return "\n".join([c for c in chunks if c is not None]).strip()
        except Exception as e:
            raise Exception(f"Error extracting PDF file {file_path}: {str(e)}")
